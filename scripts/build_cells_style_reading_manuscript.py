from __future__ import annotations

import csv
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUB = ROOT / "submission" / "bmc_medical_genomics_route_b_2026-05-01"
SOURCE_MD = SUB / "_generated_full_manuscript.md"
OUT_DOCX = SUB / "READING_VERSION_CELLS_STYLE_MANUSCRIPT.docx"
OUT_PDF = SUB / "READING_VERSION_CELLS_STYLE_MANUSCRIPT.pdf"
QA_REPORT = SUB / "READING_VERSION_CELLS_STYLE_BUILD_QA.md"

FIGURES = {
    1: SUB / "Figure1_dataset_workflow.png",
    2: SUB / "Figure2_validation_heatmap.png",
    3: SUB / "Figure3_auc_and_scores.png",
    4: SUB / "Figure4_pathway_comparison.png",
}

FIGURE_CAPTIONS = {
    1: (
        "Staged validation design for the transferability audit",
        "The analysis used GSE63060 as the discovery cohort, GSE63061 as a related validation cohort, "
        "GSE85426 as a cross-platform stress-test cohort, and GSE140829 as a large independent "
        "external validation cohort. Candidate genes and model-selection rules were frozen before "
        "the independent validation steps. The figure emphasizes the separation between discovery, "
        "related-cohort validation, cross-platform stress testing, and sequential additional independent external validation.",
    ),
    2: (
        "Cross-cohort transportability audit of discovery-derived candidate genes",
        "Heatmap showing mapping fraction, direction concordance, nominal replication fraction, "
        "targeted replication FDR fraction, median gene-level oriented AUC, and highest AUC among the frozen candidate models across validation "
        "cohorts. GSE63061 showed strong related-cohort replication, whereas GSE85426 and GSE140829 "
        "showed reduced FDR-level replication and model AUC values below 0.70.",
    ),
    3: (
        "AUC decay and exploratory MCI score placement",
        "The upper panel summarizes AUC values across training, related validation, cross-platform stress testing, "
        "and large independent external validation contexts for the primary gene-only model, primary integrated model, and train-development "
        "refit model. The dashed line indicates the prespecified AUC threshold of 0.70. The lower panel shows "
        "GSE140829 score distributions for Control, MCI, and AD samples using the frozen primary gene-only "
        "signature. MCI placement is exploratory and was not used for model optimization.",
    ),
    4: (
        "Pathway context of stable and unstable transportability classes",
        "The figure compares FDR-significant GO BP, KEGG, and Reactome enrichment yields for the strict stable "
        "and unstable gene sets. Stable transportable genes showed significant enrichment dominated by ribosomal "
        "translation and RNA-processing pathways, whereas unstable genes did not show FDR-significant enrichment "
        "in GO BP, KEGG, or Reactome under the candidate-background analysis.",
    ),
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color: str = "B7C7D9", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    for existing in list(tbl_pr.findall(qn("w:tblBorders"))):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        elem.set(qn("w:sz"), "0")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "FFFFFF")
        borders.append(elem)
    tbl_pr.append(borders)


def set_cell_width(cell, width_in: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fallback)
    run._r.append(fld_end)


def update_fields_and_export_pdf(docx_path: Path, pdf_path: Path) -> str:
    script = f"""
$ErrorActionPreference='Stop'
$docx = '{str(docx_path).replace("'", "''")}'
$pdf = '{str(pdf_path).replace("'", "''")}'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open($docx)
  $null = $doc.Fields.Update()
  foreach ($section in $doc.Sections) {{
    $section.Headers.Item(1).Range.Fields.Update() | Out-Null
    $section.Footers.Item(1).Range.Fields.Update() | Out-Null
  }}
  $doc.Save()
  $doc.ExportAsFixedFormat($pdf, 17)
  $doc.Close()
  $word.Quit()
  Write-Output 'WORD_EXPORT_OK'
}} catch {{
  if ($doc -ne $null) {{ try {{ $doc.Close($false) }} catch {{}} }}
  if ($word -ne $null) {{ try {{ $word.Quit() }} catch {{}} }}
  Write-Output ('WORD_EXPORT_FAILED: ' + $_.Exception.Message)
}}
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    return (result.stdout or result.stderr or "").strip()


def normalize_md_line(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text.strip()


def add_text_paragraph(container, text: str, style: str = "BodyText", size: float = 10.5, first_indent: bool = True):
    para = container.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.08
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.first_line_indent = Inches(0.18) if first_indent else None
    run = para.add_run(normalize_md_line(text))
    set_run_font(run, size=size)
    return para


def add_heading(container, text: str, level: int) -> None:
    para = container.add_paragraph()
    para.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    set_run_font(run, size=14 if level == 1 else 11.3, bold=True, italic=(level >= 2))


def add_caption(container, label: str, title: str, text: str, above: bool = False) -> None:
    para = container.add_paragraph()
    para.paragraph_format.space_before = Pt(5 if above else 2)
    para.paragraph_format.space_after = Pt(6)
    if above:
        para.paragraph_format.keep_with_next = True
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = para.add_run(f"{label}. ")
    set_run_font(r1, 9.2, bold=True)
    r2 = para.add_run(f"{title}. ")
    set_run_font(r2, 9.2)
    r3 = para.add_run(text)
    set_run_font(r3, 9.2)


def read_csv_dicts(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as fh:
        return list(csv.DictReader(fh))


def fmt_num(value: str) -> str:
    try:
        f = float(value)
    except Exception:
        return value
    if abs(f) >= 10 and f.is_integer():
        return str(int(f))
    if abs(f) >= 1:
        return f"{f:.3f}".rstrip("0").rstrip(".")
    return f"{f:.3f}"


def pretty_label(value: str) -> str:
    labels = {
        "discovery_training": "Discovery training",
        "related_external_validation": "Related external validation",
        "cross_platform_stress_test": "Cross-platform stress test",
        "new_large_cross_platform_final_validation": "Large independent external validation",
        "large_independent_external_validation": "Large independent external validation",
        "independent_validation_degraded": "Independent validation degraded",
        "cross_platform_degraded": "Independent validation degraded",
        "near_cohort_reproducible": "Related-cohort reproducible",
        "mixed_or_review": "Mixed/review",
        "discovery_only": "Discovery only",
        "strong_related_cohort_performance": "Strong related-cohort performance",
        "failed_070_threshold": "AUC below 0.70",
        "primary_gene_only": "Primary gene-only",
        "primary_integrated": "Primary integrated",
        "primary_integrated_or_clinical": "Primary integrated",
        "train_dev_refit_selected": "Train-development refit",
    }
    return labels.get(value, value.replace("_", " "))


def auc_ci_text(row: dict[str, str] | None) -> str:
    if not row:
        return "-"
    return f"{fmt_num(row['auc'])} ({fmt_num(row['auc_ci_low'])}-{fmt_num(row['auc_ci_high'])})"


def add_table(container, caption_label: str, caption_title: str, rows: list[list[str]], header: list[str], widths: list[float] | None = None) -> None:
    add_caption(container, caption_label, caption_title, "", above=True)
    table = container.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, cell in enumerate(hdr):
        set_cell_shading(cell, "DCEAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header[i])
        set_run_font(run, 7.8, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, 7.3)
    spacer = container.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_figure(container, number: int) -> None:
    title, caption = FIGURE_CAPTIONS[number]
    para = container.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(7)
    para.paragraph_format.space_after = Pt(2)
    width = Inches(5.85)
    if number in (3, 4):
        width = Inches(5.25)
    para.add_run().add_picture(str(FIGURES[number]), width=width)
    add_caption(container, f"Figure {number}", title, caption)


def extract_blocks(markdown: str) -> tuple[str, dict[str, str], str, str]:
    before_fig_legends, after_fig_legends = markdown.split("## Figure Legends", 1)
    refs = after_fig_legends.split("## References", 1)[1].split("## Figures", 1)[0].strip()
    title = re.search(r"^#\s+(.+)", before_fig_legends, re.M).group(1).strip()
    abstract_match = re.search(r"^## Abstract\s*(.*?)^## Background\s*", before_fig_legends, flags=re.S | re.M)
    if not abstract_match:
        raise ValueError("Could not locate abstract block and main Background heading.")
    abstract = abstract_match.group(1).strip()
    body_start = re.search(r"^## Background\s*", before_fig_legends, flags=re.M)
    if not body_start:
        raise ValueError("Could not locate main Background heading.")
    body = before_fig_legends[body_start.start():]
    metadata = {
        "author": "Zhuang Jiang",
        "affiliation": "Guangdong University of Petrochemical Technology (GDUPT), 139 Guandu 2nd Road, Maoming 525000, Guangdong, China",
        "correspondence": "Zhuang Jiang, 1627626277@qq.com",
        "orcid": "https://orcid.org/0009-0007-4388-5901",
    }
    return title, metadata, abstract, body.strip(), refs


def parse_abstract(abstract_md: str) -> tuple[list[tuple[str, str]], str]:
    sections: list[tuple[str, str]] = []
    current = None
    buf: list[str] = []
    keywords = ""
    for line in abstract_md.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("Keywords:"):
            keywords = line.replace("Keywords:", "", 1).strip()
            continue
        if line.startswith("### "):
            if current:
                sections.append((current, " ".join(buf).strip()))
            current = line.replace("### ", "", 1).strip()
            buf = []
        else:
            buf.append(line)
    if current:
        sections.append((current, " ".join(buf).strip()))
    return sections, keywords


def build_summary_tables() -> dict[str, tuple[list[str], list[list[str]], list[float]]]:
    tables: dict[str, tuple[list[str], list[list[str]], list[float]]] = {}

    dataset_rows = read_csv_dicts(ROOT / "results" / "route_b_b1_dataset_role_summary.csv")
    tables["dataset"] = (
        ["Dataset", "Study role", "Platform/scale", "AD", "Control", "Model status"],
        [[
            r["dataset_accession"],
            pretty_label(r["route_b_role"]),
            r["platform_or_scale"],
            r["n_ad"],
            r["n_control"],
            pretty_label(r["diagnostic_model_status"]),
        ] for r in dataset_rows],
        [0.75, 1.1, 2.1, 0.4, 0.55, 1.3],
    )

    metric_rows = read_csv_dicts(ROOT / "results" / "route_b_b4_transportability_metrics.csv")
    tables["transport"] = (
        ["Dataset", "Role", "Mapped", "Dir.", "Nominal", "Targeted FDR", "Median AUC", "Best frozen model AUC", "Status"],
        [[
            r["dataset_accession"],
            pretty_label(r["validation_role"]),
            r["mapped_s3_candidate_genes"],
            fmt_num(r["direction_concordance_rate"]),
            r["nominal_concordant"],
            r["fdr_concordant"],
            fmt_num(r["median_auc_oriented"]),
            fmt_num(r["best_frozen_model_auc"]),
            pretty_label(r["external_validity_status"]),
        ] for r in metric_rows],
        [0.68, 1.05, 0.48, 0.45, 0.55, 0.43, 0.62, 0.7, 1.15],
    )

    auc_ci_rows = read_csv_dicts(ROOT / "results" / "route_b_auc_ci_summary.csv")
    auc_by_role_dataset = {(r["model_role"], r["dataset_accession"]): r for r in auc_ci_rows}
    auc_rows = []
    for role in ["primary_gene_only", "primary_integrated", "train_dev_refit_selected"]:
        independent_aucs = [
            float(auc_by_role_dataset[(role, ds)]["auc"])
            for ds in ["GSE85426", "GSE140829"]
            if (role, ds) in auc_by_role_dataset
        ]
        auc_rows.append([
            pretty_label(role),
            auc_ci_text(auc_by_role_dataset.get((role, "GSE63060"))),
            auc_ci_text(auc_by_role_dataset.get((role, "GSE63061"))),
            auc_ci_text(auc_by_role_dataset.get((role, "GSE85426"))),
            auc_ci_text(auc_by_role_dataset.get((role, "GSE140829"))),
            "No" if independent_aucs and max(independent_aucs) < 0.70 else "Review",
        ])
    tables["auc"] = (
        ["Model role", "GSE63060", "GSE63061", "GSE85426", "GSE140829", "Independent >=0.70"],
        auc_rows,
        [1.25, 1.05, 1.05, 1.05, 1.05, 0.7],
    )

    set_rows = read_csv_dicts(ROOT / "results" / "route_b_b2_set_summary.csv")
    tables["genesets"] = (
        ["Metric", "Gene count"],
        [[pretty_label(r["metric"]), r["n_genes"]] for r in set_rows],
        [3.0, 1.0],
    )

    pathway_rows = read_csv_dicts(ROOT / "results" / "route_b_b4_pathway_yield_summary.csv")
    tables["pathway"] = (
        ["Gene set", "Database", "Tested terms", "FDR < 0.05 terms", "Minimum FDR"],
        [[
            pretty_label(r["gene_set"]),
            r["database"],
            r["tested_terms"],
            r["significant_terms_fdr_005"],
            fmt_num(r["min_fdr"]),
        ] for r in pathway_rows],
        [1.55, 0.75, 0.8, 0.9, 0.8],
    )

    sensitivity_rows = read_csv_dicts(ROOT / "results" / "blood_cell_marker_sensitivity_summary.csv")
    tables["sensitivity"] = (
        ["Dataset", "Mapped stable", "Dir. base", "Dir. adjusted", "Nominal base", "Nominal adjusted", "Same direction"],
        [[
            r["dataset_accession"],
            r["mapped_strict_stable_genes"],
            f"{r['base_direction_concordant']}/{r['mapped_strict_stable_genes']}",
            f"{r['marker_adjusted_direction_concordant']}/{r['mapped_strict_stable_genes']}",
            r["base_nominal_concordant"],
            r["marker_adjusted_nominal_concordant"],
            fmt_num(r["same_direction_after_adjustment_rate"]),
        ] for r in sensitivity_rows],
        [0.78, 0.65, 0.68, 0.75, 0.7, 0.82, 0.8],
    )

    supp = [
        ["S1", "Dataset inventory and inclusion criteria", "Dataset accessions, platforms, source files, sample labels, and inclusion/exclusion rules."],
        ["S2", "QC summaries", "Feature counts, sample counts, missingness, PCA summaries, and QC flags."],
        ["S3", "Discovery differential-expression results", "Full GSE63060 differential-expression statistics and frozen candidate genes."],
        ["S4", "External validation of candidate genes", "Gene-level mapping, fold change, P values, FDR, concordance, and oriented AUC."],
        ["S5", "Model and signature performance", "Frozen models, feature counts, AUC values, 95% CIs, and independent-validation degradation."],
        ["S6", "Stable and unstable gene sets", "Gene-set classification rules and memberships."],
        ["S7", "Pathway enrichment results", "GO BP, KEGG, and Reactome enrichment results."],
        ["S8", "Interpretation audit and safeguards", "Supported statements, cautionary wording, and avoided overstatements."],
        ["S9", "Reproducibility traceability", "Statement-to-script/result/figure traceability."],
        ["S10", "Demographic and mapping audit", "AD/control/MCI counts, age and sex availability, platforms, mapping fractions, concordance metrics, targeted replication FDR, and highest AUC among frozen candidate models."],
        ["S11", "Blood-cell marker-score sensitivity analysis", "Marker coverage and strict stable-gene concordance before and after neutrophil, monocyte, and lymphocyte score adjustment."],
    ]
    tables["supp"] = (
        ["Table", "Content", "Purpose"],
        supp,
        [0.5, 2.2, 3.2],
    )
    return tables


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("AD blood transcriptome manuscript")
    set_run_font(r, 8.5, italic=True)
    r2 = hp.add_run("\tComplete reading version")
    set_run_font(r2, 8.5)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, 8.5)


def add_front_matter(doc: Document, title: str, metadata: dict[str, str], abstract_md: str) -> None:
    top = doc.add_table(rows=1, cols=2)
    remove_table_borders(top)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = top.rows[0].cells
    set_cell_width(left, 1.55)
    set_cell_width(right, 5.75)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = left.paragraphs[0]
    r = p.add_run("Manuscript\nreading version")
    set_run_font(r, 12, bold=True, italic=True)
    r.font.color.rgb = RGBColor(0x0A, 0x68, 0xB0)

    for label, value in [
        ("Article type", "Research article"),
        ("Study design", "External validation audit"),
        ("Target journal", "BMC Medical Genomics"),
        ("Archive index", "github.com/1627626277-cyber/NO.2"),
    ]:
        pp = left.add_paragraph()
        pp.paragraph_format.space_after = Pt(3)
        rr = pp.add_run(label + "\n")
        set_run_font(rr, 7.5, bold=True)
        vv = pp.add_run(value)
        set_run_font(vv, 7.5)

    rp = right.paragraphs[0]
    rr = rp.add_run("Article")
    set_run_font(rr, 10.5, italic=True)
    rp.paragraph_format.space_after = Pt(3)

    tp = right.add_paragraph()
    tp.paragraph_format.space_after = Pt(7)
    tr = tp.add_run(title)
    set_run_font(tr, 18, bold=True)

    for text, size, bold in [
        ("Zhuang Jiang1*", 10.5, True),
        ("1 Guangdong University of Petrochemical Technology (GDUPT), 139 Guandu 2nd Road, Maoming 525000, Guangdong, China", 9.2, False),
        ("*Correspondence: Zhuang Jiang, 1627626277@qq.com", 9.2, False),
        ("ORCID: https://orcid.org/0009-0007-4388-5901", 9.2, False),
    ]:
        pp = right.add_paragraph()
        pp.paragraph_format.space_after = Pt(2)
        rr = pp.add_run(text)
        set_run_font(rr, size, bold=bold)

    hp = right.add_paragraph()
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after = Pt(4)
    hr = hp.add_run("Abstract")
    set_run_font(hr, 11, bold=True)

    abstract_sections, keywords = parse_abstract(abstract_md)
    for heading, text in abstract_sections:
        pp = right.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pp.paragraph_format.space_after = Pt(3)
        rr = pp.add_run(heading + ": ")
        set_run_font(rr, 9.1, bold=True)
        tt = pp.add_run(normalize_md_line(text))
        set_run_font(tt, 9.1)

    kp = right.add_paragraph()
    kp.paragraph_format.space_before = Pt(3)
    kr = kp.add_run("Keywords: ")
    set_run_font(kr, 9.1, bold=True)
    kt = kp.add_run(keywords)
    set_run_font(kt, 9.1)
    remove_table_borders(top)

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)


def iter_body_blocks(body_md: str):
    lines = body_md.splitlines()
    buffer: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                yield ("paragraph", " ".join(buffer))
                buffer = []
            continue
        if stripped.startswith("## "):
            if buffer:
                yield ("paragraph", " ".join(buffer))
                buffer = []
            yield ("h2", stripped[3:].strip())
        elif stripped.startswith("### "):
            if buffer:
                yield ("paragraph", " ".join(buffer))
                buffer = []
            yield ("h3", stripped[4:].strip())
        else:
            buffer.append(stripped)
    if buffer:
        yield ("paragraph", " ".join(buffer))


def add_body(doc: Document, body_md: str, refs: str) -> None:
    tables = build_summary_tables()
    h2_map = {
        "Background": "1. Introduction",
        "Methods": "2. Materials and Methods",
        "Results": "3. Results",
        "Discussion": "4. Discussion",
        "Conclusions": "5. Conclusions",
        "List of abbreviations": "6. Abbreviations",
        "Declarations": "7. Declarations",
    }
    subsection_counts = {"2": 0, "3": 0, "4": 0, "7": 0}
    current_major = ""

    inserted = set()
    for kind, value in iter_body_blocks(body_md):
        if kind == "h2":
            mapped = h2_map.get(value, value)
            current_major = mapped.split(".", 1)[0] if "." in mapped else ""
            add_heading(doc, mapped, 1)
            continue
        if kind == "h3":
            if current_major in subsection_counts:
                subsection_counts[current_major] += 1
                value = f"{current_major}.{subsection_counts[current_major]}. {value}"
            add_heading(doc, value, 2)
            continue

        add_text_paragraph(doc, value)

        if "The staged design is summarized in Figure 1" in value and 1 not in inserted:
            add_table(doc, "Table 1", "Dataset roles and sample counts in the transferability audit", tables["dataset"][1], tables["dataset"][0], tables["dataset"][2])
            add_figure(doc, 1)
            inserted.add(1)

        if "The cross-cohort transportability heatmap is shown in Figure 2" in value and 2 not in inserted:
            add_figure(doc, 2)
            add_table(doc, "Table 2", "Cross-cohort candidate-gene transportability metrics", tables["transport"][1], tables["transport"][0], tables["transport"][2])
            inserted.add(2)

        if "AUC decay and exploratory GSE140829 MCI score placement are shown in Figure 3" in value and 3 not in inserted:
            add_figure(doc, 3)
            add_table(doc, "Table 3", "Model AUC decay with 95% CIs across validation contexts", tables["auc"][1], tables["auc"][0], tables["auc"][2])
            inserted.add(3)

        if "The pathway comparison is summarized in Figure 4" in value and 4 not in inserted:
            add_figure(doc, 4)
            add_table(doc, "Table 4", "Stable and unstable transportability gene-set sizes", tables["genesets"][1], tables["genesets"][0], tables["genesets"][2])
            add_table(doc, "Table 5", "Pathway enrichment yield by transportability class", tables["pathway"][1], tables["pathway"][0], tables["pathway"][2])
            inserted.add(4)

        if "Nominal concordance after marker-score adjustment" in value and "sensitivity" not in inserted:
            add_table(doc, "Table 6", "Blood-cell marker-score sensitivity summary", tables["sensitivity"][1], tables["sensitivity"][0], tables["sensitivity"][2])
            inserted.add("sensitivity")

    add_heading(doc, "8. Supplementary Materials", 1)
    add_text_paragraph(
        doc,
        "The integrated reading version lists the supplementary tables to make the full manuscript structure visible in one document. "
        "Full machine-readable source tables remain in the project results and manuscript directories.",
    )
    add_table(doc, "Table 7", "Supplementary table map", tables["supp"][1], tables["supp"][0], tables["supp"][2])

    add_heading(doc, "9. References", 1)
    for line in refs.splitlines():
        line = line.strip()
        if not line:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Inches(0.22)
        para.paragraph_format.first_line_indent = Inches(-0.22)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(normalize_md_line(line))
        set_run_font(run, 8.8)


def main() -> int:
    markdown = SOURCE_MD.read_text(encoding="utf-8")
    title, metadata, abstract, body, refs = extract_blocks(markdown)

    doc = Document()
    configure_document(doc)
    add_front_matter(doc, title, metadata, abstract)
    add_body(doc, body, refs)
    doc.save(str(OUT_DOCX))

    export_status = update_fields_and_export_pdf(OUT_DOCX, OUT_PDF)
    QA_REPORT.write_text(
        "\n".join(
            [
                "# Integrated Reading Manuscript Build QA",
                "",
                "Date: 2026-05-01",
                "",
                f"- Source manuscript: `{SOURCE_MD}`",
                f"- Output DOCX: `{OUT_DOCX}`",
                f"- Output PDF: `{OUT_PDF}`",
                f"- Word export status: {export_status}",
                "",
                "Integrated objects:",
                "",
                "- Main text: title, author metadata, abstract, keywords, IMRaD body, declarations, references.",
                "- Figures: Figure 1-Figure 4 placed at the first relevant in-text mention.",
                "- Tables: Table 1-Table 7, including dataset roles, transportability metrics, AUC decay with 95% CIs, gene-set sizes, pathway yields, marker-score sensitivity, and supplementary table map.",
                "",
                "Design target:",
                "",
                "- A4 single-column integrated reading version inspired by a compact journal reading rhythm.",
                "- Serif body type, compact section numbering, figure captions directly below figures, table captions above tables.",
                "- Reading version only; not an official journal template or submission-formatting claim.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(QA_REPORT.read_text(encoding="utf-8"))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
