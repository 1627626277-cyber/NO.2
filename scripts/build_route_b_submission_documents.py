from __future__ import annotations

from pathlib import Path
import re
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUB = ROOT / "submission" / "bmc_medical_genomics_route_b_2026-05-01"
PANDOC = Path(r"C:\Users\jz\anaconda3\Library\bin\pandoc.exe")
EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")


def run(cmd: list[str]) -> None:
    print("+", " ".join(str(c) for c in cmd))
    subprocess.run(cmd, cwd=ROOT, check=True)


def slug_file_uri(path: Path) -> str:
    return path.resolve().as_uri()


def manuscript_with_figures() -> str:
    manuscript = (ROOT / "manuscript" / "ROUTE_B_MANUSCRIPT_DRAFT_V1_3_BMC_STYLE.md").read_text(encoding="utf-8")
    figure_block = f"""

## Figures

![Figure 1. Staged validation design for the transferability audit.]({(SUB / 'Figure1_dataset_workflow.png').as_posix()}){{width=6.5in}}

![Figure 2. Cross-cohort transportability audit of discovery-derived candidate genes.]({(SUB / 'Figure2_validation_heatmap.png').as_posix()}){{width=6.5in}}

![Figure 3. AUC decay and exploratory MCI score placement.]({(SUB / 'Figure3_auc_and_scores.png').as_posix()}){{width=6.5in}}

![Figure 4. Pathway context of stable and unstable transportability classes.]({(SUB / 'Figure4_pathway_comparison.png').as_posix()}){{width=6.5in}}
"""
    return manuscript.rstrip() + "\n" + figure_block


def abstract_word_count(text: str) -> int:
    match = re.search(r"^## Abstract(.*?)^## Background", text, flags=re.S | re.M)
    if not match:
        return 0
    abstract = re.sub(r"#+\s+\w+", " ", match.group(1))
    return len(re.findall(r"\S+", abstract))


def main_word_count(text: str) -> int:
    before_refs = text.split("## References", 1)[0]
    return len(re.findall(r"\S+", before_refs))


def write_html_css(path: Path) -> None:
    css = """
body {
  font-family: Arial, Helvetica, sans-serif;
  color: #111827;
  line-height: 1.45;
  max-width: 7.1in;
  margin: 0.55in auto;
  font-size: 10.5pt;
}
h1 {
  font-size: 18pt;
  line-height: 1.2;
  margin: 0 0 14pt 0;
}
h2 {
  font-size: 14pt;
  margin: 18pt 0 8pt 0;
  border-bottom: 1px solid #d1d5db;
  padding-bottom: 3pt;
}
h3 {
  font-size: 11.5pt;
  margin: 12pt 0 5pt 0;
}
p {
  margin: 0 0 8pt 0;
}
img {
  max-width: 100%;
  height: auto;
  display: block;
  margin: 10pt auto 16pt auto;
}
table {
  border-collapse: collapse;
  width: 100%;
  font-size: 9pt;
}
td, th {
  border: 1px solid #d1d5db;
  padding: 5pt;
  vertical-align: top;
}
@page {
  size: A4;
  margin: 0.65in;
}
"""
    path.write_text(css.strip() + "\n", encoding="utf-8")


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fallback)
    run._r.append(fld_end)


def style_run_font(run, name: str, size_pt: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def enable_line_numbering(section) -> None:
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:lnNumType"))):
        sect_pr.remove(existing)
    ln_num = OxmlElement("w:lnNumType")
    ln_num.set(qn("w:countBy"), "1")
    ln_num.set(qn("w:start"), "1")
    ln_num.set(qn("w:distance"), "360")
    ln_num.set(qn("w:restart"), "continuous")
    sect_pr.append(ln_num)


def update_docx_fields_with_word(docx_path: Path) -> str:
    escaped = str(docx_path).replace("'", "''")
    script = f"""
$ErrorActionPreference='Stop'
$docx = '{escaped}'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open($docx)
  $doc.PageSetup.LineNumbering.Active = $true
  $doc.PageSetup.LineNumbering.StartingNumber = 1
  $doc.PageSetup.LineNumbering.CountBy = 1
  $doc.PageSetup.LineNumbering.RestartMode = 0
  try {{ $doc.PageSetup.LineNumbering.DistanceFromText = 18 }} catch {{}}
  $null = $doc.Fields.Update()
  foreach ($section in $doc.Sections) {{
    $section.Footers.Item(1).Range.Fields.Update() | Out-Null
  }}
  $doc.Save()
  $doc.Close()
  $word.Quit()
  Write-Output 'WORD_COM_UPDATE_OK'
}} catch {{
  if ($doc -ne $null) {{ try {{ $doc.Close($false) }} catch {{}} }}
  if ($word -ne $null) {{ try {{ $word.Quit() }} catch {{}} }}
  Write-Output ('WORD_COM_UPDATE_SKIPPED: ' + $_.Exception.Message)
}}
"""
    completed = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )
    status = (completed.stdout or completed.stderr or "").strip()
    if completed.returncode != 0:
        status = f"WORD_COM_UPDATE_FAILED_EXIT_{completed.returncode}: {status}"
    print(status)
    return status


def write_submission_ready_docx(source_docx: Path, target_docx: Path) -> None:
    doc = Document(str(source_docx))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    for style_name, size in [
        ("Title", 16),
        ("Heading 1", 14),
        ("Heading 2", 12),
        ("Heading 3", 11),
    ]:
        if style_name in styles:
            styles[style_name].font.name = "Arial"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            styles[style_name].font.size = Pt(size)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        enable_line_numbering(section)
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(paragraph)
        for run in paragraph.runs:
            style_run_font(run, "Arial", 10)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            style_run_font(run, "Arial")

    for table in doc.tables:
        table.allow_autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        style_run_font(run, "Arial", 9)

    target_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target_docx))


def main() -> int:
    SUB.mkdir(parents=True, exist_ok=True)
    full_md = SUB / "_generated_full_manuscript.md"
    cover_md = SUB / "COVER_LETTER_DRAFT.md"
    css = SUB / "manuscript_print.css"
    full_docx = SUB / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_ROUTE_B_FULL.docx"
    submission_docx = SUB / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_ROUTE_B_SUBMISSION_READY.docx"
    full_html = SUB / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_ROUTE_B_FULL.html"
    full_pdf = SUB / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_ROUTE_B_FULL.pdf"
    cover_docx = SUB / "COVER_LETTER_BMC_MEDICAL_GENOMICS_ROUTE_B.docx"
    cover_html = SUB / "COVER_LETTER_BMC_MEDICAL_GENOMICS_ROUTE_B.html"
    cover_pdf = SUB / "COVER_LETTER_BMC_MEDICAL_GENOMICS_ROUTE_B.pdf"

    text = manuscript_with_figures()
    full_md.write_text(text, encoding="utf-8")
    write_html_css(css)

    if not PANDOC.exists():
        raise FileNotFoundError(PANDOC)
    if not EDGE.exists():
        raise FileNotFoundError(EDGE)

    resource_path = str(ROOT) + ";" + str(SUB)
    run([str(PANDOC), str(full_md), "--resource-path", resource_path, "-o", str(full_docx)])
    write_submission_ready_docx(full_docx, submission_docx)
    word_update_status = update_docx_fields_with_word(submission_docx)
    run([str(PANDOC), str(full_md), "--resource-path", resource_path, "--standalone", "--css", str(css), "-o", str(full_html)])
    run([
        str(EDGE),
        "--headless",
        "--disable-gpu",
        "--no-sandbox",
        f"--print-to-pdf={full_pdf}",
        slug_file_uri(full_html),
    ])

    run([str(PANDOC), str(cover_md), "-o", str(cover_docx)])
    run([str(PANDOC), str(cover_md), "--standalone", "--css", str(css), "-o", str(cover_html)])
    run([
        str(EDGE),
        "--headless",
        "--disable-gpu",
        "--no-sandbox",
        f"--print-to-pdf={cover_pdf}",
        slug_file_uri(cover_html),
    ])

    summary = SUB / "BUILD_OUTPUT_SUMMARY.md"
    summary.write_text(
        "\n".join([
            "# Build Output Summary",
            "",
            "Generated by `scripts/build_route_b_submission_documents.py`.",
            "",
            f"- Manuscript Markdown: `{full_md}`",
            f"- Manuscript DOCX: `{full_docx}`",
            f"- Submission-ready DOCX: `{submission_docx}`",
            f"- Submission-ready DOCX Word field update: {word_update_status}",
            f"- Manuscript HTML: `{full_html}`",
            f"- Manuscript PDF: `{full_pdf}`",
            f"- Cover letter DOCX: `{cover_docx}`",
            f"- Cover letter HTML: `{cover_html}`",
            f"- Cover letter PDF: `{cover_pdf}`",
            "",
            f"- Main manuscript word count before references and figures: {main_word_count(text)}",
            f"- Abstract word count: {abstract_word_count(text)}",
        ]) + "\n",
        encoding="utf-8",
    )
    print(summary.read_text(encoding="utf-8"))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
